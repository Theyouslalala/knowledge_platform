"""Document loading and parsing."""

import asyncio
from dataclasses import dataclass
from pathlib import Path


@dataclass
class ProcessedDocument:
    content: str
    metadata: dict
    file_type: str


class DocumentProcessor:
    SUPPORTED_TYPES = {".txt", ".md", ".pdf", ".docx"}

    async def process(self, file_path: str) -> ProcessedDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {suffix}")

        def _read_and_stat():
            if suffix in (".txt", ".md"):
                c = path.read_text(encoding="utf-8")
            elif suffix == ".pdf":
                c = self._read_pdf(path)
            elif suffix == ".docx":
                c = self._read_docx(path)
            else:
                c = path.read_text(encoding="utf-8")
            return c, path.stat()

        content, stat = await asyncio.to_thread(_read_and_stat)
        return ProcessedDocument(
            content=content,
            metadata={
                "filename": path.name,
                "file_type": suffix,
                "file_size": stat.st_size,
            },
            file_type=suffix,
        )

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except Exception as e:
            return f"Error reading PDF: {e}"

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            from docx import Document

            doc = Document(str(path))
            return "\n\n".join(
                p.text for p in doc.paragraphs if p.text.strip()
            )
        except Exception as e:
            return f"Error reading DOCX: {e}"
